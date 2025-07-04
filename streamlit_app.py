import streamlit as st
import streamlit.components.v1 as components
from docx import Document
import re
import os
import time
import base64
import sqlite3
import uuid

st.set_page_config(page_title="القوانين اليمنية بآخر تعديلاتها حتى عام 2025م", layout="wide")
st.markdown("<h1 style='text-align: center;'>مرحبًا بك في تطبيق القوانين اليمنية بآخر تعديلاتها حتى عام 2025م</h1>", unsafe_allow_html=True)

TRIAL_DURATION = 300 # 5 minutes in seconds (يجب أن تتطابق مع لوحة التحكم)
DATABASE_FILE = os.path.join(os.path.dirname(__file__), "user_data.db")

def init_db():
    """تهيئة قاعدة البيانات وإنشاء الجداول إذا لم تكن موجودة."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    # جدول المستخدمين
    c.execute('''
        CREATE TABLE IF NOT EXISTS users (
            user_id TEXT PRIMARY KEY,
            is_activated INTEGER DEFAULT 0,
            trial_start_time REAL,
            last_activity_time REAL,
            activation_code_used TEXT
        )
    ''')
    # جدول أكواد التفعيل
    c.execute('''
        CREATE TABLE IF NOT EXISTS activation_codes (
            code TEXT PRIMARY KEY,
            is_used INTEGER DEFAULT 0,
            used_by_user_id TEXT,
            FOREIGN KEY (used_by_user_id) REFERENCES users(user_id)
        )
    ''')
    # جدول طلبات التفعيل الجديدة
    c.execute('''
        CREATE TABLE IF NOT EXISTS activation_requests (
            request_id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            activation_code TEXT,
            request_time REAL NOT NULL,
            status TEXT DEFAULT 'pending', -- 'pending', 'approved', 'rejected'
            FOREIGN KEY (user_id) REFERENCES users(user_id),
            FOREIGN KEY (activation_code) REFERENCES activation_codes(code)
        )
    ''')
    conn.commit()
    conn.close()

def get_user_id():
    user_id_file = "user_id.txt"
    if os.path.exists(user_id_file):
        with open(user_id_file, "r") as f:
            st.session_state.user_id = f.read().strip()
    else:
        new_id = str(uuid.uuid4())
        st.session_state.user_id = new_id
        with open(user_id_file, "w") as f:
            f.write(new_id)

    # تأكد من إدخال المستخدم في قاعدة البيانات
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("INSERT OR IGNORE INTO users (user_id) VALUES (?)", (st.session_state.user_id,))
    conn.commit()
    conn.close()

    return st.session_state.user_id

def update_last_activity(user_id):
    """تحديث وقت آخر نشاط للمستخدم."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("UPDATE users SET last_activity_time = ? WHERE user_id = ?", (time.time(), user_id))
    conn.commit()
    conn.close()

def is_activated(user_id):
    """التحقق مما إذا كان المستخدم مفعلًا من قاعدة البيانات."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("SELECT is_activated FROM users WHERE user_id = ?", (user_id,))
    result = c.fetchone()
    conn.close()
    return result[0] == 1 if result else False

def get_trial_start_time(user_id):
    """الحصول على وقت بدء تجربة المستخدم من قاعدة البيانات."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("SELECT trial_start_time FROM users WHERE user_id = ?", (user_id,))
    result = c.fetchone()
    conn.close()
    return result[0] if result else None

def set_trial_start_time(user_id):
    """تعيين وقت بدء تجربة المستخدم في قاعدة البيانات."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("INSERT OR IGNORE INTO users (user_id) VALUES (?)", (user_id,)) # للتأكد من وجود المستخدم
    c.execute("UPDATE users SET trial_start_time = ?, is_activated = 0 WHERE user_id = ?", (time.time(), user_id))
    conn.commit()
    conn.close()

def send_activation_request(user_id, code):
    """إرسال طلب تفعيل للتطبيق من قبل المستخدم."""
    # --- بداية عبارات DEBUG print ---
    print(f"DEBUG: send_activation_request called for user_id: {user_id}, code: {code}") 
    # --- نهاية عبارات DEBUG print ---
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()

    # التحقق من صلاحية الكود قبل إرسال الطلب
    c.execute("SELECT is_used FROM activation_codes WHERE code = ?", (code,))
    code_status = c.fetchone()
    # --- بداية عبارات DEBUG print ---
    print(f"DEBUG: Code status for '{code}': {code_status}") 
    # --- نهاية عبارات DEBUG print ---

    if code_status and code_status[0] == 0: # الكود موجود وغير مستخدم
        # --- بداية عبارات DEBUG print ---
        print(f"DEBUG: Code '{code}' is valid and unused. Proceeding to send request.") 
        # --- نهاية عبارات DEBUG print ---
        try:
            # قم بوضع الكود في حالة "قيد الاستخدام المؤقت" لمنع مستخدم آخر من طلبه
            c.execute("UPDATE activation_codes SET is_used = 1, used_by_user_id = ? WHERE code = ?", (user_id, code))
            # --- بداية عبارات DEBUG print ---
            print(f"DEBUG: Updated activation_codes table. Rows affected: {c.rowcount}") 
            # --- نهاية عبارات DEBUG print ---
            
            # إنشاء طلب تفعيل جديد
            request_id = str(uuid.uuid4())
            c.execute("INSERT INTO activation_requests (request_id, user_id, activation_code, request_time, status) VALUES (?, ?, ?, ?, 'pending')",
                      (request_id, user_id, code, time.time()))
            # --- بداية عبارات DEBUG print ---
            print(f"DEBUG: Inserted into activation_requests. Rows affected: {c.rowcount}, request_id: {request_id}") 
            # --- نهاية عبارات DEBUG print ---
            conn.commit()
            # --- بداية عبارات DEBUG print ---
            print("DEBUG: Database commit successful.") 
            # --- نهاية عبارات DEBUG print ---
            conn.close()
            return True
        except Exception as e:
            # --- بداية عبارات DEBUG print ---
            print(f"DEBUG ERROR: An exception occurred during request sending: {e}") 
            # --- نهاية عبارات DEBUG print ---
            st.error(f"حدث خطأ أثناء إرسال الطلب: {e}")
            conn.rollback()
            conn.close()
            return False
    else:
        # --- بداية عبارات DEBUG print ---
        print(f"DEBUG: Code '{code}' is either not found or already used. Code status: {code_status}") 
        # --- نهاية عبارات DEBUG print ---
        conn.close()
        return False

def get_activation_request_status(user_id):
    """الحصول على حالة طلب التفعيل للمستخدم."""
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("SELECT status FROM activation_requests WHERE user_id = ? ORDER BY request_time DESC LIMIT 1", (user_id,))
    result = c.fetchone()
    conn.close()
    return result[0] if result else None

def highlight_keywords(text, keywords):
    text = str(text)
    text = text.replace('\xa0', ' ').replace('\u200b', '') 
    
    for kw in keywords:
        text = re.sub(f"({re.escape(kw)})", r"<mark>\1</mark>", text, flags=re.IGNORECASE | re.UNICODE)
    return text

def extract_context(paragraphs, keywords, context_lines=3): 
    paragraphs = [str(p).replace('\xa0', ' ').replace('\u200b', '') for p in paragraphs] 
    
    search_pattern = re.compile('|'.join([re.escape(kw) for kw in keywords]), re.IGNORECASE | re.UNICODE)
    
    matched_indexes = []
    for i, line in enumerate(paragraphs):
        if search_pattern.search(line):
            matched_indexes.append(i)
            
    context_set = set()
    for idx in matched_indexes:
        for i in range(max(0, idx - context_lines), min(len(paragraphs), idx + context_lines + 1)):
            context_set.add(i)
            
    filtered_paragraphs = [paragraphs[i] for i in sorted(context_set) if paragraphs[i].strip()]
    return "\n".join(filtered_paragraphs)


def export_results_to_docx(results, filename="نتائج_البحث.docx"):
    doc = Document()
    doc.add_heading("نتائج البحث", 0)
    for r in results:
        doc.add_heading(f'{r["law"]} - المادة {r["num"]}', level=1)
        doc.add_paragraph(r["context"])
    filepath = os.path.join(os.getcwd(), filename)
    doc.save(filepath)
    return filepath

def run_main_app_logic():
    components.html("""
    <style>
    .scroll-btn {
        position: fixed;
        left: 10px;
        padding: 12px;
        font-size: 24px;
        border-radius: 50%;
        background-color: #c5e1a5;
        color: black;
        cursor: pointer;
        z-index: 9999;
        border: none;
        box-shadow: 1px 1px 5px #888;
    }
    #scroll-top-btn { bottom: 80px; }
    #scroll-bottom-btn { bottom: 20px; }
    </style>
    <button class='scroll-btn' id='scroll-top-btn' onclick='window.scrollTo({top: 0, behavior: "smooth"});'>⬆️</button>
    <button class='scroll-btn' id='scroll-bottom-btn' onclick='window.scrollTo({top: document.body.scrollHeight, behavior: "smooth"});'>⬇️</button>
    """, height=1)

    subfolders = [f.path for f in os.scandir() if f.is_dir() and f.name not in [".git", ".streamlit"]]
    if not subfolders:
        st.warning("📂 لا توجد مجلدات قوانين.")
        return

    selected_folder = st.selectbox("اختر مجلدًا للبحث فيه:", ["🔍 كل المجلدات"] + subfolders)

    all_files = {}
    if selected_folder == "🔍 كل المجلدات":
        for folder in subfolders:
            files = [f for f in os.listdir(folder) if f.endswith(".docx")]
            all_files[folder] = files
    else:
        files = [f for f in os.listdir(selected_folder) if f.endswith(".docx")]
        all_files[selected_folder] = files

    keywords = st.text_area("الكلمات المفتاحية (افصل بفاصلة)", "")

    if "results" not in st.session_state:
        st.session_state.results = []
    if "search_done" not in st.session_state:
        st.session_state.search_done = False

    if st.button("🔍 بدء البحث") and keywords:
        kw_list = [k.strip() for k in keywords.split(",") if k.strip()]
        results = []

        for folder, files in all_files.items():
            for file in files:
                doc_path = os.path.join(folder, file)
                try:
                    doc = Document(doc_path)
                except Exception as e:
                    st.warning(f"⚠️ تعذر قراءة الملف {file} في المجلد {folder}: {e}. قد يكون الملف تالفًا أو مشفرًا.")
                    continue

                law_name = file.replace(".docx", "")
                
                all_paragraphs_in_doc = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                current_article_paragraphs = []
                last_article_num = "غير معروفة"

                for i, para_text in enumerate(all_paragraphs_in_doc): 
                    match = re.match(r"مادة\s*\(?\s*(\d+)\)?", para_text)
                    if match:
                        if current_article_paragraphs:
                            full_article_text = "\n".join(current_article_paragraphs)
                            if any(kw.lower() in full_article_text.lower() for kw in kw_list):
                                context = extract_context(current_article_paragraphs, kw_list, context_lines=3) 
                                results.append({
                                    "law": law_name,
                                    "num": last_article_num,
                                    "text": highlight_keywords(context, kw_list),
                                    "plain": full_article_text,
                                    "context": context,
                                    "keywords": kw_list
                                })
                            current_article_paragraphs = []
                        last_article_num = match.group(1)
                        current_article_paragraphs.append(para_text)
                    else:
                        current_article_paragraphs.append(para_text)
                
                if current_article_paragraphs:
                    full_article_text = "\n".join(current_article_paragraphs)
                    if any(kw.lower() in full_article_text.lower() for kw in kw_list):
                        context = extract_context(current_article_paragraphs, kw_list, context_lines=3) 
                        results.append({
                            "law": law_name,
                            "num": last_article_num,
                            "text": highlight_keywords(context, kw_list),
                            "plain": full_article_text,
                            "context": context,
                            "keywords": kw_list
                        })

        st.session_state.results = results
        st.session_state.search_done = True

    if st.session_state.search_done and st.session_state.results:
        results = st.session_state.results
        unique_laws = sorted(set(r["law"] for r in results))
        st.success(f"تم العثور على {len(results)} نتيجة في {len(unique_laws)} قانون/ملف.")
        
        selected_law = st.selectbox("فلترة حسب القانون", ["الكل"] + unique_laws)
        filtered = results if selected_law == "الكل" else [r for r in results if r["law"] == selected_law]

        for r in filtered:
            st.markdown(f"""
<div style="background-color:#f1f8e9;padding:15px;margin-bottom:15px;border-radius:10px;
            border:1px solid #c5e1a5;direction:rtl;text-align:right; overflow-wrap: break-word;">
    <p style="font-weight:bold;font-size:18px;margin:0">🔷 {r["law"]} - المادة {r["num"]}</p>
    <p style="font-size:17px;line-height:1.8;margin-top:10px">
        {r["text"]}
    </p>
</div>
""", unsafe_allow_html=True)

        if filtered:
            filepath = export_results_to_docx(filtered)
            with open(filepath, "rb") as f:
                st.download_button(
                    label="📥 تحميل النتائج كملف Word",
                    data=f,
                    file_name="نتائج_البحث.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

def main():
    init_db()
    user_id = get_user_id()
    update_last_activity(user_id)

    if "activated" not in st.session_state:
        st.session_state.activated = is_activated(user_id)

    if not st.session_state.activated:
        st.warning("⚠️ التطبيق غير مفعل. يرجى التفعيل أو استخدام النسخة التجريبية.")
        
        trial_start_time = get_trial_start_time(user_id)
        
        if trial_start_time is None:
            if st.button("🕒 بدء التجربة المجانية", key="start_trial_button"):
                set_trial_start_time(user_id)
                st.session_state.trial_start_time = time.time()
                st.success("🎉 بدأت النسخة التجريبية. لديك 5 دقائق. يرجى تحديث الصفحة أو إعادة تشغيل التطبيق.")
                st.rerun() 
        else:
            time_elapsed = time.time() - trial_start_time
            if time_elapsed < TRIAL_DURATION:
                remaining_minutes = int((TRIAL_DURATION - time_elapsed) / 60)
                st.info(f"✅ النسخة التجريبية نشطة. تبقى لديك حوالي {remaining_minutes} دقيقة.")
                run_main_app_logic()
            else:
                st.error("❌ انتهت مدة التجربة المجانية. يرجى التفعيل.")
        
        # --- بداية التعديلات لآلية طلب التفعيل ---
        request_status = get_activation_request_status(user_id)
        
        if request_status == 'pending':
            st.info("⏳ لقد أرسلت طلب تفعيل. نرجو الانتظار حتى تتم موافقة المسؤول. يمكنك تحديث الصفحة للتحقق من الحالة.")
        elif request_status == 'rejected':
            st.error("❌ تم رفض طلب التفعيل الخاص بك من قبل المسؤول. يرجى التواصل مع الدعم أو محاولة استخدام كود تفعيل آخر.")
            code = st.text_input("أدخل كود التفعيل هنا", key="activation_code_input_rejected")
            if st.button("🔐 إرسال طلب تفعيل جديد", key="send_new_request_button_rejected"):
                if code and send_activation_request(user_id, code.strip()):
                    st.success("✅ تم إرسال طلب التفعيل بنجاح! سيتم مراجعته من قبل المسؤول.")
                    st.rerun()
                else:
                    st.error("❌ كود التفعيل غير صحيح أو مستخدم بالفعل أو حدث خطأ أثناء إرسال الطلب.")
        else: # لا يوجد طلب معلق أو كان مرفوضا وتم مسح الحالة
            code = st.text_input("أدخل كود التفعيل هنا", key="activation_code_input")
            if st.button("🔐 إرسال طلب تفعيل", key="send_activation_request_button"):
                if code and send_activation_request(user_id, code.strip()):
                    st.success("✅ تم إرسال طلب التفعيل بنجاح! سيتم مراجعته من قبل المسؤول.")
                    st.rerun() 
                else:
                    st.error("❌ كود التفعيل غير صحيح أو مستخدم بالفعل أو حدث خطأ أثناء إرسال الطلب. يرجى التأكد من الكود.")
        # --- نهاية التعديلات لآلية طلب التفعيل ---

    else:
        run_main_app_logic()

main()
